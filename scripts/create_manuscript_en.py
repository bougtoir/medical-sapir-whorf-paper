#!/usr/bin/env python3
"""
Generate the English manuscript for:
"The Nosological Imperative: Disease Naming as Therapeutic Intervention
 in Pain Medicine — A Medical Sapir-Whorf Perspective"

Output: medical_sapir_whorf/output/manuscript_en.docx
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def add_superscript_text(paragraph, text):
    """Parse text with {N} or {N-M} markers and render as superscript runs."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(9)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
    return paragraph


def set_cell_text(cell, text, bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Set text in a table cell with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def create_manuscript():
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)

    # ══════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Article Type: Perspective")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()  # blank line

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "The Nosological Imperative: Disease Naming as Therapeutic Intervention "
        "in Pain Medicine — A Medical Sapir-Whorf Perspective"
    )
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Author names to be added]")
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Affiliations
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Affiliations to be added]")
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Corresponding author
    p = doc.add_paragraph()
    run = p.add_run("Corresponding Author: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run("[Name, address, email]")
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Word count
    p = doc.add_paragraph()
    run = p.add_run("Word count: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run("approximately 4,500 (excluding references)")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("Number of tables: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run("2")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("Number of figures: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run("1")
    run.font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════
    # ABSTRACT
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ABSTRACT")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    abstract_text = (
        "The Sapir-Whorf hypothesis posits that language shapes cognition and perception. "
        "We propose that an analogous mechanism operates in medicine: nosological categories — "
        "the names and classifications of diseases — do not merely describe clinical reality but "
        "actively constitute it. We term this the Medical Sapir-Whorf Hypothesis. While the "
        "general principle that classification systems influence medical practice may seem obvious, "
        "we argue that the effects extend far beyond administrative convenience into genuinely "
        "therapeutic territory. Pain medicine provides a compelling domain in which to examine "
        "this thesis. Chronic pain has been historically under-classified, and patients with "
        "unclassifiable pain syndromes experience not only physical suffering but also the "
        "existential distress of diagnostic uncertainty. The introduction of a dedicated chronic "
        "pain chapter (MG30) in the International Classification of Diseases, 11th Revision "
        "(ICD-11), the formalization of nociplastic pain as a third mechanistic descriptor by "
        "the International Association for the Study of Pain (IASP), and the establishment of "
        "chronic primary pain as a disease entity in its own right collectively represent a "
        "natural experiment in therapeutic naming. We review evidence that diagnostic labeling "
        "in pain medicine influences patient beliefs, treatment-seeking behavior, clinician "
        "management decisions, research funding allocation, and health policy. We propose a "
        "theoretical framework — the Nosological Therapeutic Loop — in which disease naming "
        "triggers a self-reinforcing cycle of patient validation, clinical attention, research "
        "investment, and treatment development. We outline a research agenda to empirically "
        "test this framework using ICD-10 to ICD-11 transition data, and discuss the boundary "
        "conditions that distinguish therapeutic naming from disease mongering."
    )
    p = doc.add_paragraph()
    run = p.add_run(abstract_text)
    run.font.size = Pt(11)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Keywords: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        "Sapir-Whorf hypothesis; nosology; chronic pain; ICD-11; diagnostic labeling; "
        "nociplastic pain; pain classification; therapeutic naming"
    )
    run.font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════
    # 1. INTRODUCTION
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("1. INTRODUCTION")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    intro_paras = [
        (
            "In 1929, Edward Sapir observed that \"the 'real world' is to a large extent "
            "unconsciously built up on the language habits of the group.\"{1} His student "
            "Benjamin Lee Whorf extended this insight into the hypothesis that linguistic "
            "categories shape thought and perception — a principle now known as the "
            "Sapir-Whorf hypothesis or linguistic relativity.{2} While the strong version "
            "of this hypothesis (linguistic determinism) has been largely abandoned, the "
            "weaker version — that language influences cognition without fully determining "
            "it — has accumulated substantial empirical support across domains including "
            "color perception, spatial reasoning, and temporal cognition.{3}"
        ),
        (
            "We propose that an analogous mechanism operates in clinical medicine. "
            "Nosological categories — the names, definitions, and classification systems "
            "through which diseases are organized — do not merely label pre-existing "
            "clinical entities but actively shape how clinicians perceive symptoms, how "
            "patients experience illness, how researchers allocate effort, and how health "
            "systems distribute resources. We term this the Medical Sapir-Whorf Hypothesis."
        ),
        (
            "At first glance, this may appear trivially true: medicine is a social "
            "institution, and of course its conceptual frameworks influence its practice. "
            "However, we argue that the effects of nosological categories extend into "
            "genuinely therapeutic territory — that the act of naming a disease can itself "
            "constitute a form of treatment. If this is correct, then the field of pain "
            "medicine, where chronic conditions have been historically under-classified "
            "and patients frequently endure years of diagnostic uncertainty, represents "
            "the domain where nosological expansion carries the greatest therapeutic potential."
        ),
        (
            "This perspective article develops the Medical Sapir-Whorf Hypothesis with "
            "specific reference to pain medicine. We draw on three recent developments — "
            "the ICD-11 chronic pain classification (MG30), the formalization of nociplastic "
            "pain by the International Association for the Study of Pain (IASP), and "
            "emerging evidence on the psychological effects of diagnostic labeling — to "
            "argue that pain medicine is uniquely positioned to benefit from deliberate "
            "nosological expansion. We propose a formal theoretical model (the Nosological "
            "Therapeutic Loop), outline a research agenda, and address the critical "
            "boundary between therapeutic naming and disease mongering."
        ),
    ]

    for text in intro_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 2. THEORETICAL FRAMEWORK
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("2. THEORETICAL FRAMEWORK: FROM LINGUISTIC TO NOSOLOGICAL RELATIVITY")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("2.1 The Sapir-Whorf Hypothesis in Its Original Form")
    run.bold = True
    run.font.size = Pt(11)

    framework_paras = [
        (
            "The classical Sapir-Whorf hypothesis exists along a continuum from strong "
            "determinism (language determines thought) to weak relativism (language influences "
            "thought). Contemporary research supports the weaker form: speakers of languages "
            "with different color term boundaries show measurably different performance on "
            "color discrimination tasks,{3} and speakers of languages with absolute spatial "
            "reference frames demonstrate superior dead-reckoning abilities compared to "
            "speakers of languages using relative frames.{4}"
        ),
        (
            "Critically, these effects are not merely epiphenomenal. Linguistic categories "
            "create cognitive scaffolding that facilitates certain patterns of attention, "
            "memory, and reasoning while inhibiting others. The categories do not create "
            "the phenomena they describe — colors exist independently of color terms — but "
            "they shape how those phenomena are perceived, processed, and acted upon."
        ),
    ]
    for text in framework_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("2.2 The Medical Sapir-Whorf Hypothesis")
    run.bold = True
    run.font.size = Pt(11)

    medical_sw_paras = [
        (
            "We define the Medical Sapir-Whorf Hypothesis as follows: nosological categories "
            "shape clinical cognition, patient experience, and healthcare system behavior in "
            "ways that extend beyond the mere administrative function of classification. "
            "Specifically, we propose that disease names operate through at least five "
            "distinct channels:"
        ),
    ]
    for text in medical_sw_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    # Numbered list of channels
    channels = [
        (
            "Perceptual channeling: ",
            "Disease categories direct clinician attention toward certain symptom "
            "constellations and away from others. A clinician who has the concept of "
            "\"nociplastic pain\" will perceive and interpret a patient's presentation "
            "differently from one who lacks this category.{5}"
        ),
        (
            "Diagnostic foreclosure: ",
            "Once a disease name is assigned, clinical investigation tends to stop "
            "at the boundary of that category. The label becomes a cognitive endpoint "
            "that both enables treatment and constrains further inquiry.{6}"
        ),
        (
            "Patient ontological validation: ",
            "Receiving a disease name transforms a patient's subjective experience from "
            "\"symptoms without explanation\" to \"a recognized medical condition.\" This "
            "transition carries measurable psychological and even physiological consequences.{7}"
        ),
        (
            "Research infrastructure alignment: ",
            "Funding bodies, specialty journals, and clinical trial registries are organized "
            "around disease categories. Conditions without recognized nosological status "
            "are systematically under-researched — not because they are unimportant, but "
            "because they lack the categorical infrastructure necessary to attract "
            "sustained scientific attention.{8}"
        ),
        (
            "Policy and resource allocation: ",
            "Health systems allocate resources — specialist training positions, dedicated "
            "clinics, reimbursement codes — according to recognized disease categories. "
            "Unnamed conditions are structurally invisible to healthcare planning.{9}"
        ),
    ]

    for i, (bold_part, normal_part) in enumerate(channels, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(f"({i}) ")
        run.font.size = Pt(11)
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(11)
        add_superscript_text(p, normal_part)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("2.3 Precedents: Hacking's Looping Effects and Kleinman's Illness Narratives")
    run.bold = True
    run.font.size = Pt(11)

    precedent_paras = [
        (
            "Our framework builds on two important intellectual precedents. Ian Hacking's "
            "concept of \"looping effects\" describes how classificatory categories in the "
            "human sciences interact with the people they classify: the category shapes "
            "behavior, the changed behavior feeds back into the category, and the category "
            "evolves in response.{10} In psychiatry, this dynamic is well documented — DSM "
            "categories influence patient self-identification, which in turn shapes symptom "
            "presentation, which eventually prompts diagnostic revision.{11}"
        ),
        (
            "Arthur Kleinman's distinction between disease (the biomedical entity) and "
            "illness (the lived experience) is equally relevant.{12} Kleinman showed that "
            "clinical encounters are sites of negotiation between explanatory models, and "
            "that the physician's act of naming — translating illness into disease — carries "
            "profound consequences for the patient's experience. Notably, Kleinman also "
            "documented that certain illness experiences, such as anorexia nervosa, persist "
            "even in cultural contexts where the prevailing social values would seem to "
            "preclude them, suggesting that while naming shapes experience, it does not "
            "create disease de novo.{12}"
        ),
        (
            "The Medical Sapir-Whorf Hypothesis synthesizes these insights: disease names "
            "are neither passive labels for pre-existing natural kinds nor arbitrary social "
            "constructions, but active cognitive-institutional tools that shape the clinical "
            "reality they purport to describe."
        ),
    ]
    for text in precedent_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 3. PAIN MEDICINE AS THE IDEAL TEST CASE
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("3. PAIN MEDICINE AS THE IDEAL TEST CASE")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("3.1 The Historical Under-Classification of Chronic Pain")
    run.bold = True
    run.font.size = Pt(11)

    pain_history_paras = [
        (
            "Chronic pain affects approximately one in five adults globally and represents "
            "a leading cause of years lived with disability.{13,14} Despite this enormous "
            "burden, chronic pain was poorly represented in the International Classification "
            "of Diseases (ICD) through its tenth revision. In ICD-10, chronic pain diagnoses "
            "were scattered across multiple chapters (F, G, M, R), did not differentiate "
            "between acute and chronic presentations, and included ambiguous categories such "
            "as \"R52.2 Other chronic pain\" that provided no meaningful clinical information.{15}"
        ),
        (
            "The consequences of this nosological poverty were not merely administrative. "
            "A study comparing ICD-10 and ICD-11 coding in a Thai pain clinic found that "
            "under ICD-10, roughly one-third of patient encounters were coded as \"other "
            "chronic pain,\" and the next most common codes specified pain region rather than "
            "underlying cause.{16} Chronic pain was, in a precise nosological sense, invisible: "
            "it could not be counted, compared, or prioritized because the language to describe "
            "it did not exist within the classification system."
        ),
    ]
    for text in pain_history_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("3.2 The ICD-11 Chronic Pain Classification: A Nosological Revolution")
    run.bold = True
    run.font.size = Pt(11)

    icd11_paras = [
        (
            "The ICD-11, approved by the World Health Assembly in 2019, introduced a dedicated "
            "chapter for chronic pain under MG30, developed by an international task force "
            "convened by IASP.{17} This classification established seven main categories "
            "(Table 1), each with multiple diagnostic levels ranging from primary care to "
            "specialized pain treatment."
        ),
    ]
    for text in icd11_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ── TABLE 1: ICD-11 Chronic Pain Classification ──
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    run = p.add_run("Table 1. ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "ICD-11 Chronic Pain Classification (MG30) with Medical Sapir-Whorf implications"
    )
    run.font.size = Pt(10)

    table1_data = [
        ["ICD-11 Code", "Category", "Key Innovation over ICD-10", "Sapir-Whorf Implication"],
        [
            "MG30.0",
            "Chronic primary pain",
            "Pain as disease entity, not symptom",
            "Validates pain without identifiable cause; reduces diagnostic nihilism",
        ],
        [
            "MG30.1",
            "Chronic cancer-related pain",
            "Distinguishes cancer pain from cancer itself",
            "Enables pain-specific treatment alongside oncology",
        ],
        [
            "MG30.2",
            "Chronic postsurgical or post-traumatic pain",
            "Previously invisible category",
            "Makes iatrogenic/injury-related chronification visible",
        ],
        [
            "MG30.3",
            "Chronic secondary musculoskeletal pain",
            "Differentiates chronic from acute MSK pain",
            "Shifts framing from structural pathology to pain chronification",
        ],
        [
            "MG30.4",
            "Chronic secondary visceral pain",
            "Previously coded as organ-specific diagnoses only",
            "Recognizes pain as independent clinical target",
        ],
        [
            "MG30.5",
            "Chronic neuropathic pain",
            "Unified category across etiologies",
            "Creates coherent research and treatment framework",
        ],
        [
            "MG30.6",
            "Chronic secondary headache or orofacial pain",
            "Chronic-specific classification",
            "Distinguishes chronic management from acute episodes",
        ],
    ]

    table1 = doc.add_table(rows=len(table1_data), cols=4)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(table1_data):
        row = table1.rows[i]
        for j, cell_text in enumerate(row_data):
            bold = (i == 0)
            set_cell_text(row.cells[j], cell_text, bold=bold, size=9)

    doc.add_paragraph()

    icd11_cont_paras = [
        (
            "The most conceptually radical innovation is MG30.0, Chronic primary pain, which "
            "recognizes chronic pain as a disease in its own right — not merely a symptom of "
            "an underlying condition.{17} This single nosological act exemplifies the Medical "
            "Sapir-Whorf Hypothesis: by creating a name for pain-without-identifiable-cause, "
            "ICD-11 transforms the clinical encounter. The patient who previously received "
            "the implicit message \"we cannot find what is wrong with you\" now receives \"you "
            "have chronic primary pain\" — a recognized, codable, treatable condition."
        ),
        (
            "A WHO field test demonstrated that ICD-11 achieved 63.2% correct chronic pain "
            "diagnoses compared to 43.0% with ICD-10, with the improvement most dramatic for "
            "cases where chronic pain was regarded as a symptom of an underlying disease "
            "(63.5% vs. 26.8%).{18} This improvement reflects not just better coding but "
            "better clinical thinking: the new categories provide cognitive scaffolding that "
            "helps clinicians perceive and classify pain presentations more accurately."
        ),
    ]
    for text in icd11_cont_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("3.3 Nociplastic Pain: Naming the Previously Unnameable")
    run.bold = True
    run.font.size = Pt(11)

    nociplastic_paras = [
        (
            "The 2017 introduction of \"nociplastic pain\" by IASP as a third mechanistic "
            "descriptor — alongside nociceptive and neuropathic pain — provides a further "
            "illustration of therapeutic naming.{19} Before this term existed, patients whose "
            "pain could not be attributed to either tissue damage (nociceptive) or nerve "
            "lesion (neuropathic) occupied a nosological void. Conditions such as fibromyalgia, "
            "irritable bowel syndrome, and chronic tension-type headache lacked a unified "
            "mechanistic framework."
        ),
        (
            "The term \"nociplastic\" — referring to pain arising from altered nociception "
            "despite no clear evidence of actual or threatened tissue damage or disease of "
            "the somatosensory nervous system — fills this void.{19} It provides clinicians "
            "with a positive diagnosis rather than a diagnosis of exclusion; it gives "
            "researchers a coherent category around which to organize investigation of central "
            "sensitization mechanisms; and it gives patients a name that validates their "
            "experience without implying psychogenic causation.{20}"
        ),
        (
            "The clinical impact of this naming act is already becoming apparent. A Lancet "
            "review noted that the concept of nociplastic pain has enabled a more nuanced "
            "understanding of prevalent pain conditions and has begun to reshape treatment "
            "approaches toward mechanism-based rather than disease-based therapeutics.{21}"
        ),
    ]
    for text in nociplastic_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 4. EVIDENCE FOR THERAPEUTIC EFFECTS OF NAMING
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("4. EVIDENCE FOR THE THERAPEUTIC EFFECTS OF DIAGNOSTIC NAMING IN PAIN")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("4.1 Diagnostic Labeling and Patient Beliefs")
    run.bold = True
    run.font.size = Pt(11)

    labeling_paras = [
        (
            "A growing body of experimental evidence demonstrates that diagnostic labels "
            "exert measurable effects on patient beliefs and behavior. O'Keeffe et al. "
            "conducted a six-arm randomized experiment (n = 1,375) examining the effect of "
            "different labels for low back pain on patient beliefs about imaging, surgery, "
            "seriousness, and recovery.{22} They found that pathoanatomical labels (\"disc "
            "bulge,\" \"degeneration,\" \"arthritis\") led to significantly higher perceived need "
            "for imaging and surgery, lower recovery expectations, and greater perceived "
            "seriousness compared to non-specific labels (\"episode of back pain,\" \"lumbar "
            "sprain\"). These differences were amplified among patients with current pain "
            "who had a history of seeking care."
        ),
        (
            "A systematic review of randomized trials on diagnostic labeling in "
            "musculoskeletal pain confirmed these findings, reporting that specific diagnostic "
            "labels may increase patient preferences for invasive care and foster more "
            "negative recovery expectations, while non-specific labels led to more positive "
            "recovery beliefs but lower patient satisfaction.{23} This satisfaction paradox "
            "is itself evidence for the Medical Sapir-Whorf Hypothesis: patients desire "
            "specific names for their conditions even when those names carry prognostic "
            "disadvantages, suggesting that the naming function serves psychological needs "
            "beyond information transmission."
        ),
    ]
    for text in labeling_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("4.2 Diagnostic Uncertainty and Patient Suffering")
    run.bold = True
    run.font.size = Pt(11)

    uncertainty_paras = [
        (
            "The converse of diagnostic naming — diagnostic uncertainty — provides further "
            "evidence for the therapeutic valence of nosological categories. A grounded theory "
            "study of patients with uncertain diagnoses found that intolerance of uncertainty "
            "generated substantial psychological distress independent of physical symptoms, "
            "including heightened anxiety, reduced self-efficacy, and impaired coping.{24} "
            "In chronic pain specifically, diagnostic uncertainty has been associated with "
            "greater functional disability, more catastrophic thinking, and increased "
            "healthcare utilization.{25}"
        ),
        (
            "A survey of 100 chronic pain clinic patients found that all participants valued "
            "receiving an accurate diagnosis, with 76% strongly agreeing. Diagnostic clarity "
            "was reported to improve psychological well-being, confidence in treatment, and "
            "overall satisfaction with care.{26} These findings suggest that the act of "
            "receiving a diagnosis — being given a name for one's suffering — constitutes a "
            "therapeutic intervention in itself, separate from any specific treatment that "
            "may follow."
        ),
    ]
    for text in uncertainty_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("4.3 Clinician Behavior and the Nosological Framing Effect")
    run.bold = True
    run.font.size = Pt(11)

    clinician_paras = [
        (
            "The effects of disease naming extend beyond patients to clinicians themselves. "
            "The availability of a diagnostic category influences clinical decision-making "
            "through what we term the Nosological Framing Effect. When chronic pain is "
            "classified as \"R52.2 Other chronic pain\" (ICD-10), the implicit clinical "
            "message is that no further diagnostic refinement is possible or necessary. When "
            "the same presentation is classified as \"MG30.0 Chronic primary pain\" (ICD-11), "
            "the clinician receives a positive diagnostic framework with associated treatment "
            "guidelines, referral pathways, and outcome measures."
        ),
        (
            "This framing effect operates at the system level as well. The introduction of "
            "specific chronic pain codes in ICD-11 has the potential to reshape referral "
            "patterns, specialist training curricula, and quality improvement metrics. When "
            "a condition can be named and coded, it can be measured; when it can be measured, "
            "it can be managed; when it can be managed, resources can be allocated to it. "
            "The nosological act creates the institutional infrastructure for care."
        ),
    ]
    for text in clinician_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 5. THE NOSOLOGICAL THERAPEUTIC LOOP
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("5. THE NOSOLOGICAL THERAPEUTIC LOOP: A FORMAL MODEL")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    loop_paras = [
        (
            "Building on the evidence reviewed above, we propose a formal model — the "
            "Nosological Therapeutic Loop (NTL) — describing the self-reinforcing cycle "
            "through which disease naming generates therapeutic effects (Fig. 1). The NTL "
            "consists of six interconnected stages:"
        ),
    ]
    for text in loop_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    # Stages of the loop
    stages = [
        (
            "Nosological Act (N): ",
            "A disease category is created or refined within a classification system "
            "(e.g., ICD-11 MG30.0 Chronic primary pain)."
        ),
        (
            "Patient Validation (V): ",
            "Patients receive a recognized name for their condition, reducing diagnostic "
            "uncertainty and providing ontological legitimacy to their suffering."
        ),
        (
            "Clinical Attention (A): ",
            "Clinicians acquire new cognitive categories that direct attention, shape "
            "differential diagnosis, and guide treatment selection."
        ),
        (
            "Research Investment (R): ",
            "The named condition attracts funding, dedicated journals, specialty conferences, "
            "and clinical trials organized around the new category."
        ),
        (
            "Treatment Development (T): ",
            "Research generates new therapeutic approaches specifically tailored to the "
            "named condition."
        ),
        (
            "Nosological Refinement (N\u2032): ",
            "Clinical experience and research findings prompt revision and refinement of "
            "the original category, initiating a new cycle."
        ),
    ]

    for i, (bold_part, normal_part) in enumerate(stages, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(f"Stage {i}: ")
        run.font.size = Pt(11)
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(normal_part)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # Figure 1 inline image
    fig1_path = os.path.join(OUTPUT_DIR, "figure1_ntl_en.png")
    if os.path.exists(fig1_path):
        p = doc.add_paragraph()
        p.space_before = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig1_path, width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.space_before = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[Figure 1: Run create_ntl_figure.py first to generate the image]")
        run.font.size = Pt(10)
        run.italic = True

    p = doc.add_paragraph()
    p.space_before = Pt(6)
    run = p.add_run("Figure 1. ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "The Nosological Therapeutic Loop (NTL). Disease naming (N) triggers a self-reinforcing "
        "cycle through patient validation (V), clinical attention (A), research investment (R), "
        "treatment development (T), and nosological refinement (N\u2032). Each stage feeds into the "
        "next, creating a positive feedback loop. The red circle indicates the critical boundary "
        "condition: if the loop operates without empirical grounding in patient outcomes, it "
        "risks degenerating into disease mongering."
    )
    run.font.size = Pt(10)

    doc.add_paragraph()

    ntl_cont_paras = [
        (
            "The NTL model is explicitly analogous to Hacking's looping effects but extends "
            "it in two important ways. First, it identifies the therapeutic valence of the "
            "loop: the cycle is not merely descriptive (categories shape the classified) but "
            "normatively beneficial (the shaping improves patient outcomes). Second, it "
            "specifies the institutional mechanisms (funding, training, coding) through which "
            "the loop operates, moving beyond the individual-level focus of Hacking's original "
            "formulation.{10}"
        ),
        (
            "The NTL can be expressed semi-formally. Let D denote a disease category, P the "
            "patient population, C the clinical community, and O patient outcomes. The Medical "
            "Sapir-Whorf Hypothesis predicts that \u2202O/\u2202D > 0 when D appropriately captures "
            "a genuine pattern of suffering — that is, the introduction of a well-founded "
            "disease category improves patient outcomes through the mechanisms described above, "
            "holding treatment technology constant."
        ),
    ]
    for text in ntl_cont_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 6. BOUNDARY CONDITIONS
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run(
        "6. BOUNDARY CONDITIONS: DISTINGUISHING THERAPEUTIC NAMING FROM DISEASE MONGERING"
    )
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    boundary_paras = [
        (
            "The Medical Sapir-Whorf Hypothesis raises an obvious concern: if naming diseases "
            "is therapeutic, does this not justify unlimited nosological expansion? The "
            "concept of \"disease mongering\" — the widening of disease boundaries to grow "
            "markets for those who sell treatments — represents a well-documented pathology "
            "of nosological excess.{27} The pharmaceutical industry's promotion of conditions "
            "such as \"social anxiety disorder\" and \"female sexual dysfunction\" has been "
            "criticized as creating diseases to match available drugs rather than developing "
            "drugs to match genuine diseases.{28}"
        ),
        (
            "We argue that the distinction between therapeutic naming and disease mongering "
            "can be drawn along three dimensions (Table 2):"
        ),
    ]
    for text in boundary_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ── TABLE 2: Therapeutic naming vs disease mongering ──
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    run = p.add_run("Table 2. ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Distinguishing therapeutic naming from disease mongering")
    run.font.size = Pt(10)

    table2_data = [
        ["Dimension", "Therapeutic Naming", "Disease Mongering"],
        [
            "Origin",
            "Driven by clinical need and patient suffering; developed through consensus of clinical experts and patient advocates",
            "Driven by commercial interest; developed through pharmaceutical marketing and sponsored disease awareness campaigns",
        ],
        [
            "Evidence base",
            "Grounded in empirical observation of distinct symptom patterns, mechanisms, or treatment responses; validated through field testing",
            "Based on medicalization of normal variation or redefinition of risk factors as diseases; limited independent validation",
        ],
        [
            "Outcome trajectory",
            "Leads to improved patient outcomes, better resource allocation, and refined understanding over time (NTL positive cycle)",
            "Leads to overdiagnosis, unnecessary treatment, iatrogenic harm, and resource misallocation",
        ],
    ]

    table2 = doc.add_table(rows=len(table2_data), cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(table2_data):
        row = table2.rows[i]
        for j, cell_text in enumerate(row_data):
            bold = (i == 0)
            set_cell_text(row.cells[j], cell_text, bold=bold, size=9)

    doc.add_paragraph()

    boundary_cont_paras = [
        (
            "The ICD-11 chronic pain classification exemplifies therapeutic naming by all "
            "three criteria: it was developed by an international clinical expert task force "
            "in response to documented clinical need,{17} it was validated through WHO field "
            "testing,{18} and early evidence suggests improved diagnostic accuracy and "
            "clinical utility.{16} The contrast with disease mongering is not merely one of "
            "degree but of kind: therapeutic naming responds to existing suffering, while "
            "disease mongering creates anxiety about previously unproblematic states."
        ),
        (
            "A further boundary condition is provided by Kleinman's observation that certain "
            "disease experiences persist independently of their cultural naming.{12} Anorexia "
            "nervosa existed in eras when thinness was not valued; chronic pain persists "
            "regardless of whether it is named. The Medical Sapir-Whorf Hypothesis does not "
            "claim that naming creates disease — it claims that naming shapes the clinical "
            "response to disease. This distinction is crucial: it means that nosological "
            "expansion is therapeutic only when it names something that already exists as "
            "a pattern of suffering, not when it pathologizes normal variation."
        ),
    ]
    for text in boundary_cont_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 7. PROPOSED RESEARCH AGENDA
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("7. PROPOSED RESEARCH AGENDA")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    agenda_paras = [
        (
            "The transition from ICD-10 to ICD-11 constitutes a natural experiment in "
            "therapeutic naming that can be leveraged to empirically test the Medical "
            "Sapir-Whorf Hypothesis. We propose the following research agenda:"
        ),
    ]
    for text in agenda_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    research_items = [
        (
            "Before-after diagnostic pattern analysis: ",
            "Using administrative claims data from healthcare systems that have adopted "
            "ICD-11, compare chronic pain diagnostic patterns before and after the transition. "
            "Primary outcomes: distribution of chronic pain diagnoses, proportion of patients "
            "receiving specific (vs. non-specific) pain diagnoses, and changes in referral "
            "patterns to pain specialists."
        ),
        (
            "Patient-reported outcome study: ",
            "Prospective cohort study comparing patient satisfaction, perceived diagnostic "
            "certainty, treatment adherence, and self-reported pain outcomes between patients "
            "diagnosed using ICD-10 vs. ICD-11 categories for equivalent clinical presentations."
        ),
        (
            "Clinician cognition experiment: ",
            "Randomized vignette-based study presenting identical pain presentations to "
            "clinicians with and without training in ICD-11 chronic pain categories. Primary "
            "outcomes: diagnostic accuracy, treatment plan quality, and clinical reasoning "
            "patterns assessed via think-aloud protocol."
        ),
        (
            "Research funding analysis: ",
            "Bibliometric analysis of chronic pain research publications and funding "
            "allocations before and after ICD-11 adoption, stratified by the new MG30 "
            "subcategories. Test prediction: categories that received new or refined ICD-11 "
            "codes will show disproportionate growth in research activity."
        ),
        (
            "Cross-national comparison: ",
            "Exploit staggered ICD-11 adoption across countries to implement a "
            "difference-in-differences design comparing chronic pain outcomes in early vs. "
            "late adopter nations."
        ),
    ]

    for i, (bold_part, normal_part) in enumerate(research_items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(f"({i}) ")
        run.font.size = Pt(11)
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(normal_part)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 8. CONCLUSION
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("8. CONCLUSION")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    conclusion_paras = [
        (
            "The Sapir-Whorf hypothesis teaches that language is not a passive mirror of "
            "reality but an active participant in its construction. The Medical Sapir-Whorf "
            "Hypothesis extends this insight to clinical medicine: nosological categories are "
            "not mere labels but cognitive-institutional tools that shape perception, direct "
            "attention, allocate resources, and ultimately influence patient outcomes."
        ),
        (
            "Pain medicine provides the most compelling domain for this thesis. The historical "
            "under-classification of chronic pain in ICD-10 was not merely an administrative "
            "inconvenience — it was a structural cause of suffering, rendering millions of "
            "patients invisible to the healthcare systems ostensibly designed to help them. "
            "The ICD-11 chronic pain classification, the concept of nociplastic pain, and the "
            "recognition of chronic primary pain as a disease entity collectively represent "
            "acts of therapeutic naming: they do not merely describe clinical reality but "
            "actively improve it."
        ),
        (
            "The Nosological Therapeutic Loop provides a formal framework for understanding "
            "how naming generates therapeutic effects and for identifying the boundary "
            "conditions that separate therapeutic naming from disease mongering. The critical "
            "distinction is not between naming and not naming, but between naming that "
            "responds to genuine patterns of suffering (grounded in clinical evidence and "
            "developed through expert consensus) and naming that creates artificial disease "
            "categories to serve commercial interests."
        ),
        (
            "If the Medical Sapir-Whorf Hypothesis is correct, then the deliberate, "
            "evidence-based expansion of nosological categories in pain medicine is not "
            "merely a bureaucratic exercise in classification — it is a therapeutic "
            "intervention. The implication is both practical and profound: in pain medicine, "
            "increasing the precision and richness of disease names may be one of the most "
            "cost-effective therapeutic strategies available."
        ),
    ]
    for text in conclusion_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # CONFLICT OF INTEREST
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("CONFLICT OF INTEREST")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("The authors declare no conflicts of interest.")
    run.font.size = Pt(11)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # FUNDING
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("FUNDING")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("[Funding information to be added]")
    run.font.size = Pt(11)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("REFERENCES")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    references = [
        "Sapir E. The status of linguistics as a science. Language. 1929;5(4):207-214.",
        "Whorf BL. Language, Thought, and Reality: Selected Writings. Cambridge, MA: MIT Press; 1956.",
        (
            "Winawer J, Witthoft N, Frank MC, Wu L, Wade AR, Boroditsky L. Russian blues reveal "
            "effects of language on color discrimination. Proc Natl Acad Sci U S A. 2007;104(19):7780-7785."
        ),
        (
            "Levinson SC. Frames of reference and Molyneux's question: cross-linguistic evidence. "
            "In: Bloom P, Peterson MA, Nadel L, Garrett MF, eds. Language and Space. Cambridge, MA: "
            "MIT Press; 1996:109-169."
        ),
        (
            "Kosek E, Cohen M, Baron R, et al. Do we need a third mechanistic descriptor for chronic "
            "pain states? Pain. 2016;157(7):1382-1386."
        ),
        (
            "Jutel A. Putting a Name to It: Diagnosis in Contemporary Society. Baltimore: "
            "Johns Hopkins University Press; 2011."
        ),
        (
            "Bontempo AC. Does labeling a patient with fibromyalgia alter health care utilization? "
            "A retrospective, matched-cohort study. J Gen Intern Med. 2021;36(5):1352-1358."
        ),
        (
            "Kaplan RM. Disease, Diagnoses, and Dollars: Facing the Ever-Expanding Market for "
            "Medical Care. New York: Copernicus Books; 2009."
        ),
        (
            "World Health Organization. International Classification of Diseases 11th Revision: "
            "The Global Standard for Diagnostic Health Information. Geneva: WHO; 2019."
        ),
        (
            "Hacking I. The looping effects of human kinds. In: Sperber D, Premack D, Premack AJ, "
            "eds. Causal Cognition: A Multidisciplinary Debate. Oxford: Clarendon Press; 1995:351-394."
        ),
        (
            "Tsou JY. Hacking on the looping effects of psychiatric classifications: what is an "
            "interactive and indifferent kind? Int Stud Philos Sci. 2007;21(3):329-344."
        ),
        (
            "Kleinman A. The Illness Narratives: Suffering, Healing, and the Human Condition. "
            "New York: Basic Books; 1988."
        ),
        (
            "Breivik H, Collett B, Ventafridda V, Cohen R, Gallacher D. Survey of chronic pain in "
            "Europe: prevalence, impact on daily life, and treatment. Eur J Pain. 2006;10(4):287-333."
        ),
        (
            "GBD 2021 Diseases and Injuries Collaborators. Global incidence, prevalence, and years "
            "lived with disability for 354 diseases and injuries. Lancet. 2024;403(10440):2162-2199."
        ),
        (
            "Rief W, Kaasa S, Jensen R, et al. The need to revise pain diagnoses in ICD-10. Pain. "
            "2010;149(2):169-170."
        ),
        (
            "Nugraha B, Gutenbrunner C, Barke A, et al. The IASP classification of chronic pain "
            "for ICD-11: functioning properties of chronic pain. Pain. 2019;160(1):88-94."
        ),
        (
            "Treede RD, Rief W, Barke A, et al. Chronic pain as a symptom or a disease: the IASP "
            "Classification of Chronic Pain for the International Classification of Diseases "
            "(ICD-11). Pain. 2019;160(1):19-27."
        ),
        (
            "Barke A, Korwisi B, Jakob R, Konstanjsek N, Rief W, Treede RD. Classification of "
            "chronic pain for the International Classification of Diseases (ICD-11): results of "
            "the 2017 international WHO field testing. Pain. 2022;163(2):e285-e292."
        ),
        (
            "Kosek E, Clauw D, Nijs J, et al. Chronic nociplastic pain affecting the "
            "musculoskeletal system: clinical criteria and grading system. Pain. 2021;162(11):2629-2634."
        ),
        (
            "Nijs J, Malfliet A, Nishigami T. Nociplastic pain and central sensitization in "
            "patients with chronic pain conditions: a terminology update for clinicians. Braz J "
            "Phys Ther. 2023;27(3):100518."
        ),
        (
            "Fitzcharles MA, Cohen SP, Clauw DJ, Littlejohn G, Usui C, Häuser W. Nociplastic pain: "
            "towards an understanding of prevalent pain conditions. Lancet. 2021;397(10289):2098-2110."
        ),
        (
            "O'Keeffe M, Ferreira GE, Harris IA, et al. Effect of diagnostic labelling on "
            "management intentions for non-specific low back pain: a randomized scenario-based "
            "experiment. Eur J Pain. 2022;26(7):1532-1545."
        ),
        (
            "Zadro JR, Ferreira GE, O'Keeffe M, et al. The effect of diagnostic labels on "
            "treatment preferences and beliefs in people with musculoskeletal pain: a systematic "
            "review of randomized trials. Arch Phys Med Rehabil. 2025 (in press)."
        ),
        (
            "Sari DM, Samuels K, Engel L, et al. How patients with an uncertain diagnosis "
            "experience intolerance of uncertainty: a grounded theory study. Psychol Res Behav "
            "Manag. 2021;14:1269-1281."
        ),
        (
            "Tanna V, Heathcote LC, Heirich MS, et al. Something else going on? Diagnostic "
            "uncertainty in children with chronic pain and their parents. Children. 2020;7(10):165."
        ),
        (
            "O'Sullivan E, Shanahan E, Costelloe L, et al. Importance of diagnosis in chronic "
            "pain: patient's perspective. Pain Stud Treat. 2026;14(1):1-12."
        ),
        (
            "Moynihan R, Heath I, Henry D. Selling sickness: the pharmaceutical industry and "
            "disease mongering. BMJ. 2002;324(7342):886-891."
        ),
        (
            "Moynihan R, Cassels A. Selling Sickness: How the World's Biggest Pharmaceutical "
            "Companies Are Turning Us All into Patients. New York: Nation Books; 2005."
        ),
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        # Reference number in superscript
        run = p.add_run(f"{i}. ")
        run.font.size = Pt(10)
        run = p.add_run(ref)
        run.font.size = Pt(10)

    # ── Save ──
    output_path = os.path.join(OUTPUT_DIR, "manuscript_en.docx")
    doc.save(output_path)
    print(f"Manuscript saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_manuscript()
