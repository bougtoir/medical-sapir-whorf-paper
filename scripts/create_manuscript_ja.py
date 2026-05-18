#!/usr/bin/env python3
"""
Generate the Japanese manuscript for:
「疾患命名の治療的意義：疼痛医学における医療版サピア＝ウォーフ仮説」

Output: medical_sapir_whorf/output/manuscript_ja.docx
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def add_superscript_text(paragraph, text):
    """Parse text with {N} or {N-M} markers and render as superscript runs."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(9)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
    return paragraph


def set_cell_text(cell, text, bold=False, size=10):
    """Set text in a table cell with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def create_manuscript():
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)

    # ══════════════════════════════════════════════
    # 表紙
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("論文種別: 展望論文 (Perspective)")
    run.font.size = Pt(11)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "疾患命名の治療的意義：疼痛医学における医療版サピア＝ウォーフ仮説"
    )
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[著者名]")
    run.font.size = Pt(11)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[所属]")
    run.font.size = Pt(11)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("責任著者: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run("[氏名、住所、メール]")
    run.font.size = Pt(11)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("語数: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run("約4,500語（参考文献除く）")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("表の数: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run("2")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("図の数: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run("1")
    run.font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════
    # 抄録
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("抄録")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    abstract_text = (
        "サピア＝ウォーフ仮説は、言語が認知と知覚を形作ると主張する。本論文では、"
        "同様のメカニズムが臨床医学において作動することを提唱する。すなわち、疾患命名体系"
        "（nosological categories）は臨床的現実を単に記述するのではなく、能動的にそれを構成"
        "する。我々はこれを「医療版サピア＝ウォーフ仮説」と名付ける。分類体系が医療実践に"
        "影響を与えるという一般原則は自明に思われるかもしれないが、その効果は行政的便宜を"
        "はるかに超え、真に治療的な領域にまで及ぶことを論じる。疼痛医学はこの命題を検証する"
        "ための理想的な領域を提供する。慢性疼痛は歴史的に分類が不十分であり、分類不能な疼痛"
        "症候群を抱える患者は身体的苦痛のみならず、診断的不確実性による実存的苦悩をも経験"
        "する。国際疾病分類第11版（ICD-11）における慢性疼痛専用章（MG30）の導入、国際疼痛"
        "学会（IASP）による第三の疼痛機序記述子としてのnociplastic painの定式化、および"
        "慢性一次性疼痛の独立した疾患としての認定は、治療的命名の自然実験を構成する。"
        "診断ラベリングが患者の信念、受療行動、臨床家の管理判断、研究資金配分、および保健"
        "政策に影響を与えるエビデンスをレビューする。疾患命名が患者の承認、臨床的関心、"
        "研究投資、治療開発の自己強化的サイクルを誘発する理論的枠組み「疾患命名治療ループ」"
        "（Nosological Therapeutic Loop）を提唱し、ICD-10からICD-11への移行データを用いた"
        "実証的研究アジェンダを提示する。また、治療的命名と疾患喧伝（disease mongering）を"
        "区別する境界条件について論じる。"
    )
    p = doc.add_paragraph()
    run = p.add_run(abstract_text)
    run.font.size = Pt(11)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("キーワード: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(
        "サピア＝ウォーフ仮説; 疾患命名体系; 慢性疼痛; ICD-11; "
        "診断ラベリング; nociplastic pain; 疼痛分類; 治療的命名"
    )
    run.font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════
    # 1. 緒言
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("1. 緒言")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    intro_paras = [
        (
            '1929年、エドワード・サピアは「\u201c現実の世界\u201dは大部分、集団の言語習慣の上に無意識に'
            "構築されている」と観察した。{1} 彼の弟子であるベンジャミン・リー・ウォーフはこの"
            "洞察を、言語カテゴリーが思考と知覚を形作るという仮説へと拡張した。これは現在"
            "サピア＝ウォーフ仮説または言語的相対性として知られている。{2} この仮説の強い"
            "バージョン（言語決定論）は概ね放棄されたが、弱いバージョン——言語が認知に影響を"
            "与えるが完全には決定しない——は、色知覚、空間推論、時間認知を含む複数の領域で"
            "実質的な実証的支持を蓄積している。{3}"
        ),
        (
            "我々は、類似のメカニズムが臨床医学において作動することを提唱する。疾患命名体系"
            "——疾患が組織される名称、定義、分類体系——は既存の臨床的実体を単にラベリングする"
            "のではなく、臨床家が症状をどう認知するか、患者が疾患をどう経験するか、研究者が"
            "どこに努力を配分するか、医療システムがどう資源を分配するかを能動的に形作る。"
            "我々はこれを「医療版サピア＝ウォーフ仮説」と名付ける。"
        ),
        (
            "一見、これは自明に思えるかもしれない。医療は社会制度であり、概念的枠組みが実践に"
            "影響するのは当然である。しかし我々は、疾患命名体系の効果が真に治療的な領域にまで"
            "及ぶこと——すなわち疾患に名前をつける行為自体が一つの治療形態を構成しうること——を"
            "論じる。もしこれが正しければ、慢性疾患が歴史的に分類不足であり、患者がしばしば"
            "何年もの診断的不確実性に耐えている疼痛医学こそ、疾患命名体系の拡張が最大の治療的"
            "可能性を有する領域である。"
        ),
        (
            "本展望論文は、疼痛医学に特に言及しながら医療版サピア＝ウォーフ仮説を展開する。"
            "ICD-11慢性疼痛分類（MG30）、IASPによるnociplastic painの定式化、および診断"
            "ラベリングの心理的効果に関する新たなエビデンスという三つの最近の進展を活用し、"
            "疼痛医学が意図的な疾患命名の拡張から最も恩恵を受ける立場にあることを論じる。"
            "理論モデル（疾患命名治療ループ）を提唱し、研究アジェンダを示し、治療的命名と"
            "疾患喧伝の境界について論じる。"
        ),
    ]

    for text in intro_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 2. 理論的枠組み
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("2. 理論的枠組み：言語的相対性から疾患命名的相対性へ")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("2.1 サピア＝ウォーフ仮説の原形")
    run.bold = True
    run.font.size = Pt(11)

    sw_paras = [
        (
            "古典的サピア＝ウォーフ仮説は、強い決定論（言語が思考を決定する）から弱い相対論"
            "（言語が思考に影響を与える）までの連続体上に存在する。現代の研究は弱い形式を支持"
            "している。色彩語の境界が異なる言語の話者は、色弁別課題で測定可能な異なるパフォー"
            "マンスを示し、{3} 絶対的空間参照枠を持つ言語の話者は、相対的枠組みを使用する"
            "言語の話者と比較して優れた方向感覚能力を示す。{4}"
        ),
        (
            "重要なのは、これらの効果は単なる付随現象ではないということである。言語カテゴリーは"
            "注意、記憶、推論の特定のパターンを促進し、他のパターンを抑制する認知的足場を作る。"
            "カテゴリーはそれが記述する現象を創出しない——色は色彩語とは独立に存在する——が、"
            "それらの現象がいかに知覚され、処理され、行動の対象となるかを形作る。"
        ),
    ]
    for text in sw_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("2.2 医療版サピア＝ウォーフ仮説")
    run.bold = True
    run.font.size = Pt(11)

    msw_paras = [
        (
            "医療版サピア＝ウォーフ仮説を以下のように定義する。疾患命名体系は、分類という"
            "行政的機能を超えて、臨床的認知、患者経験、および医療システムの行動を形作る。"
            "具体的には、疾患名は少なくとも五つの経路を通じて作用すると提唱する："
        ),
    ]
    for text in msw_paras:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    channels_ja = [
        ("知覚的水路化: ", "疾患カテゴリーは臨床家の注意を特定の症状群へ向け、他の症状群から逸らす。{5}"),
        ("診断的閉鎖: ", "疾患名が割り当てられると、臨床的探索はそのカテゴリーの境界で停止する傾向がある。{6}"),
        ("患者の存在論的承認: ", "疾患名を受け取ることは、患者の主観的経験を「説明のない症状」から「認知された医学的状態」へと変容させる。{7}"),
        ("研究基盤の整列: ", "資金提供機関、専門誌、臨床試験登録は疾患カテゴリーに沿って組織される。認知された疾患命名上の地位を持たない状態は体系的に研究不足となる。{8}"),
        ("政策・資源配分: ", "医療システムは認知された疾患カテゴリーに従って資源を配分する。名前のない状態は医療計画に対して構造的に不可視である。{9}"),
    ]

    for i, (bold_part, normal_part) in enumerate(channels_ja, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(f"({i}) ")
        run.font.size = Pt(11)
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(11)
        add_superscript_text(p, normal_part)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("2.3 先行研究：Hackingのループ効果とKleinmanの病の語り")
    run.bold = True
    run.font.size = Pt(11)

    precedent_ja = [
        (
            "我々の枠組みは二つの重要な知的先行研究に基づく。イアン・ハッキングの「ループ効果」"
            "は、人間科学における分類カテゴリーが分類される人々と相互作用する様を記述する。"
            "カテゴリーが行動を形作り、変化した行動がカテゴリーにフィードバックし、カテゴリーが"
            "応答して進化する。{10} 精神医学においてこのダイナミクスはよく文書化されている。"
            "DSMのカテゴリーが患者の自己同定に影響し、それが症状の発現を形作り、最終的に診断の"
            "改訂を促す。{11}"
        ),
        (
            "アーサー・クラインマンの疾患（disease：生物医学的実体）と病い（illness：生きられた"
            "経験）の区別も同様に関連する。{12} クラインマンは臨床的出会いが説明モデル間の交渉"
            "の場であり、医師の命名行為——病いを疾患へと翻訳すること——が患者の経験に深い帰結を"
            "もたらすことを示した。注目すべきは、クラインマンは拒食症のような特定の病い経験が、"
            "それを排除するように思われる社会的価値が支配的な文化的文脈においても存続すること"
            "を文書化し、命名が経験を形作る一方で疾患をゼロから創出するわけではないことを示唆"
            "している。{12}"
        ),
        (
            "医療版サピア＝ウォーフ仮説はこれらの洞察を統合する。疾患名は既存の自然種に対する"
            "受動的ラベルでも恣意的な社会的構成でもなく、それが記述すると称する臨床的現実を"
            "形作る能動的な認知的・制度的ツールである。"
        ),
    ]
    for text in precedent_ja:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 3. 疼痛医学が理想的テストケースである理由
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("3. 疼痛医学が理想的テストケースである理由")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("3.1 慢性疼痛の歴史的分類不足")
    run.bold = True
    run.font.size = Pt(11)

    pain_ja = [
        (
            "慢性疼痛は世界の成人の約5人に1人に影響を及ぼし、障害生存年数の主要な原因を構成"
            "する。{13,14} この膨大な疾病負担にもかかわらず、慢性疼痛は国際疾病分類（ICD）の"
            "第10版まで不十分にしか表現されてこなかった。ICD-10では慢性疼痛の診断は複数の章"
            "（F, G, M, R）に散在し、急性と慢性を区別せず、「R52.2 その他の慢性疼痛」のような"
            "臨床的に意味のない曖昧なカテゴリーを含んでいた。{15}"
        ),
        (
            "タイのペインクリニックにおけるICD-10とICD-11コーディングの比較研究では、ICD-10の"
            "下で患者来院の約3分の1が「その他の慢性疼痛」としてコードされ、次に多いコードは原因"
            "ではなく疼痛の部位を特定するものであった。{16} 慢性疼痛は正確な疾患命名上の意味で"
            "不可視であった。それを記述する言語が分類体系内に存在しなかったため、数えることも"
            "比較することも優先順位をつけることもできなかった。"
        ),
    ]
    for text in pain_ja:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("3.2 ICD-11慢性疼痛分類：疾患命名の革命")
    run.bold = True
    run.font.size = Pt(11)

    icd11_ja = [
        (
            "2019年に世界保健総会で承認されたICD-11は、IASPが招集した国際タスクフォースにより"
            "開発されたMG30の下に慢性疼痛の専用章を導入した。{17} この分類は7つの主要カテゴリー"
            "（表1）を確立し、それぞれがプライマリケアから専門的疼痛治療に至る複数の診断レベルを"
            "有する。"
        ),
    ]
    for text in icd11_ja:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ── 表1 ──
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    run = p.add_run("表1. ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run("ICD-11慢性疼痛分類（MG30）と医療版サピア＝ウォーフ仮説の含意")
    run.font.size = Pt(10)

    table1_ja = [
        ["ICD-11コード", "カテゴリー", "ICD-10からの革新", "サピア＝ウォーフ的含意"],
        ["MG30.0", "慢性一次性疼痛", "疼痛を症状ではなく疾患として認定", "原因不明の疼痛を承認し診断的虚無主義を軽減"],
        ["MG30.1", "慢性がん関連疼痛", "がんの疼痛をがんそのものから区別", "腫瘍学と並行した疼痛特異的治療を可能に"],
        ["MG30.2", "慢性術後・外傷後疼痛", "以前は不可視のカテゴリー", "医原性・損傷関連の慢性化を可視化"],
        ["MG30.3", "慢性二次性筋骨格系疼痛", "慢性と急性の筋骨格系疼痛を区別", "構造的病理から疼痛慢性化へ枠組みを転換"],
        ["MG30.4", "慢性二次性内臓痛", "以前は臓器特異的診断のみ", "疼痛を独立した臨床的ターゲットとして認識"],
        ["MG30.5", "慢性神経障害性疼痛", "病因横断的な統一カテゴリー", "一貫した研究・治療枠組みを創出"],
        ["MG30.6", "慢性二次性頭痛・口腔顔面痛", "慢性特異的分類", "慢性管理を急性エピソードから区別"],
    ]

    table1 = doc.add_table(rows=len(table1_ja), cols=4)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(table1_ja):
        row = table1.rows[i]
        for j, cell_text in enumerate(row_data):
            set_cell_text(row.cells[j], cell_text, bold=(i == 0), size=9)

    doc.add_paragraph()

    icd11_ja_cont = [
        (
            "最も概念的に革命的な革新はMG30.0、慢性一次性疼痛であり、慢性疼痛を基礎疾患の"
            "単なる症状ではなくそれ自体の疾患として認定する。{17} この単一の疾患命名行為は"
            "医療版サピア＝ウォーフ仮説を例証する。原因不明の疼痛に名前を創ることにより、"
            "ICD-11は臨床的出会いを変容させる。以前は「あなたの不調の原因が見つかりません」"
            "という暗黙のメッセージを受けていた患者が、今や「あなたは慢性一次性疼痛です」"
            "——認知され、コード化され、治療可能な状態——を受け取る。"
        ),
        (
            "WHOのフィールドテストは、ICD-11が慢性疼痛診断の正確率63.2%を達成し、ICD-10の"
            "43.0%と比較して改善したことを示した。改善は特に慢性疼痛が基礎疾患の症状とみなされる"
            "症例で顕著であった（63.5% vs. 26.8%）。{18}"
        ),
    ]
    for text in icd11_ja_cont:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("3.3 Nociplastic Pain：名付けられなかったものに名前をつける")
    run.bold = True
    run.font.size = Pt(11)

    noci_ja = [
        (
            "2017年にIASPが侵害受容性・神経障害性に続く第三の疼痛機序記述子として"
            "「nociplastic pain」を導入したことは、治療的命名のさらなる例証である。{19} "
            "この用語が存在する前は、組織損傷（侵害受容性）にも神経病変（神経障害性）にも"
            "帰属できない疼痛を持つ患者は疾患命名の空白を占めていた。線維筋痛症、過敏性腸"
            "症候群、慢性緊張型頭痛などの状態は統一的な機序的枠組みを欠いていた。"
        ),
        (
            "「nociplastic」という用語は、この空白を埋める。{19} 除外診断ではなく積極的診断を"
            "臨床家に提供し、中枢感作メカニズムの研究を組織する一貫したカテゴリーを研究者に"
            "与え、心因性の暗示なしに経験を承認する名前を患者に与える。{20}"
        ),
        (
            "Lancetのレビューは、nociplastic painの概念が有病率の高い疼痛状態のより精緻な理解を"
            "可能にし、疾患ベースではなく機序ベースの治療アプローチへの再構築を開始したと"
            "指摘している。{21}"
        ),
    ]
    for text in noci_ja:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 4. 命名の治療的効果のエビデンス
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("4. 疼痛における診断命名の治療的効果のエビデンス")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("4.1 診断ラベリングと患者の信念")
    run.bold = True
    run.font.size = Pt(11)

    label_ja = [
        (
            "診断ラベルが患者の信念と行動に測定可能な効果を及ぼすことを示す実験的エビデンスが"
            "蓄積されている。O'Keeffeらは腰痛に対する異なるラベルの効果を検証する6群ランダム化"
            "実験（n = 1,375）を実施した。{22} 病理解剖学的ラベル（「椎間板膨隆」「変性」"
            "「関節炎」）は、非特異的ラベル（「腰痛エピソード」「腰部捻挫」）と比較して画像検査"
            "・手術の必要性の認知、重症度の認知が有意に高く、回復期待が低かった。"
        ),
        (
            "筋骨格系疼痛における診断ラベリングのランダム化試験のシステマティックレビューは、"
            "特異的診断ラベルが侵襲的治療への選好を高め、より否定的な回復期待を助長する一方、"
            "非特異的ラベルはより肯定的な回復信念をもたらすが患者満足度を低下させることを確認"
            "した。{23} この満足度のパラドックスは医療版サピア＝ウォーフ仮説のエビデンスである。"
            "患者は予後的不利益を伴う場合でも自身の状態の特異的な名前を望むのであり、命名機能が"
            "情報伝達を超えた心理的ニーズに応えていることを示唆する。"
        ),
    ]
    for text in label_ja:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("4.2 診断的不確実性と患者の苦痛")
    run.bold = True
    run.font.size = Pt(11)

    uncert_ja = [
        (
            "診断命名の対極——診断的不確実性——は、疾患命名カテゴリーの治療的価値のさらなる"
            "エビデンスを提供する。不確定な診断を持つ患者のグラウンデッド・セオリー研究は、"
            "不確実性への不耐性が身体症状とは独立した実質的な心理的苦痛を生じさせることを"
            "見出した。{24} 慢性疼痛においては特に、診断的不確実性がより大きな機能障害、"
            "破局的思考の増加、医療利用の増加と関連している。{25}"
        ),
        (
            "慢性疼痛クリニック患者100名の調査では、全参加者が正確な診断を受けることを重視し、"
            "76%が強く同意した。診断の明確さは心理的幸福感、治療への信頼、ケアへの全体的満足度を"
            "改善すると報告された。{26}"
        ),
    ]
    for text in uncert_ja:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("4.3 臨床家の行動と疾患命名のフレーミング効果")
    run.bold = True
    run.font.size = Pt(11)

    clinician_ja = [
        (
            "疾患命名の効果は患者を超えて臨床家自身にも及ぶ。診断カテゴリーの利用可能性は、"
            "我々が「疾患命名フレーミング効果」と呼ぶものを通じて臨床的意思決定に影響する。"
            "慢性疼痛が「R52.2 その他の慢性疼痛」（ICD-10）として分類される場合、暗黙の臨床的"
            "メッセージはさらなる診断的精緻化が不可能または不必要であるということである。同じ"
            "プレゼンテーションが「MG30.0 慢性一次性疼痛」（ICD-11）として分類される場合、"
            "臨床家は関連する治療ガイドライン、紹介経路、アウトカム指標を伴う積極的な診断枠組みを"
            "受け取る。"
        ),
        (
            "このフレーミング効果はシステムレベルでも作動する。ICD-11における特定の慢性疼痛コードの"
            "導入は、紹介パターン、専門医研修カリキュラム、品質改善指標を再構築する可能性を有する。"
            "状態に名前をつけてコード化できれば測定できる。測定できれば管理できる。管理できれば"
            "資源を配分できる。疾患命名行為がケアのための制度的インフラストラクチャーを創出する。"
        ),
    ]
    for text in clinician_ja:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 5. 疾患命名治療ループ
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("5. 疾患命名治療ループ：形式モデル")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    ntl_ja = [
        (
            "以上のエビデンスに基づき、疾患命名が治療的効果を生み出す自己強化的サイクルを記述する"
            "形式モデル「疾患命名治療ループ（NTL）」を提唱する（図1）。NTLは6つの相互連結した"
            "段階からなる："
        ),
    ]
    for text in ntl_ja:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    stages_ja = [
        ("疾患命名行為（N）: ", "分類体系内に疾患カテゴリーが創設または精緻化される。"),
        ("患者の承認（V）: ", "患者が自身の状態に対する認知された名前を受け取り、診断的不確実性が軽減される。"),
        ("臨床的注目（A）: ", "臨床家が注意を向け、鑑別診断を形作り、治療選択を導く新しい認知カテゴリーを獲得する。"),
        ("研究投資（R）: ", "命名された状態が資金、専門誌、学会、臨床試験を新カテゴリー周辺に引き付ける。"),
        ("治療開発（T）: ", "研究が命名された状態に特化した新しい治療アプローチを生み出す。"),
        ("疾患命名の精緻化（N\u2032）: ", "臨床経験と研究知見が元のカテゴリーの改訂と精緻化を促し、新しいサイクルが始まる。"),
    ]

    for i, (bold_part, normal_part) in enumerate(stages_ja, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(f"段階{i}: ")
        run.font.size = Pt(11)
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(normal_part)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # 図1 インライン画像
    fig1_path = os.path.join(OUTPUT_DIR, "figure1_ntl_ja.png")
    if os.path.exists(fig1_path):
        p = doc.add_paragraph()
        p.space_before = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig1_path, width=Inches(5.5))
    else:
        p = doc.add_paragraph()
        p.space_before = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[図1：create_ntl_figure.pyを先に実行してください]")
        run.font.size = Pt(10)
        run.italic = True

    p = doc.add_paragraph()
    p.space_before = Pt(6)
    run = p.add_run("図1. ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "疾患命名治療ループ（NTL）。疾患命名（N）が患者承認（V）、臨床的注目（A）、"
        "研究投資（R）、治療開発（T）、疾患命名精緻化（N\u2032）を経る自己強化サイクルを誘発する。"
        "赤い円は境界条件を示す：ループが患者アウトカムの実証的根拠なしに作動する場合、"
        "疾患喧伝に退化するリスクがある。"
    )
    run.font.size = Pt(10)

    doc.add_paragraph()

    ntl_cont_ja = [
        (
            "NTLモデルはHackingのループ効果と明示的に類似するが、二つの重要な点で拡張する。"
            "第一に、ループの治療的価値を同定する。サイクルは単に記述的（カテゴリーが被分類者を"
            "形作る）ではなく規範的に有益（形作りが患者アウトカムを改善する）である。第二に、"
            "ループが作動する制度的メカニズム（資金、研修、コーディング）を特定し、Hackingの"
            "元の定式化の個人レベルの焦点を超える。{10}"
        ),
        (
            "NTLは半形式的に表現できる。Dを疾患カテゴリー、Pを患者集団、Cを臨床コミュニティ、"
            "Oを患者アウトカムとする。医療版サピア＝ウォーフ仮説は、Dが苦痛の真のパターンを"
            "適切に捉える場合に\u2202O/\u2202D > 0を予測する。すなわち、十分に根拠のある疾患カテゴリーの"
            "導入は、治療技術を一定として、上述のメカニズムを通じて患者アウトカムを改善する。"
        ),
    ]
    for text in ntl_cont_ja:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 6. 境界条件
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("6. 境界条件：治療的命名と疾患喧伝の区別")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    boundary_ja = [
        (
            "医療版サピア＝ウォーフ仮説は明白な懸念を提起する。疾患に名前をつけることが治療的で"
            "あるならば、無制限の疾患命名の拡張を正当化するのではないか。「疾患喧伝」——治療を"
            "販売する者のための市場を拡大するために疾患の境界を広げること——は、疾患命名の過剰の"
            "よく文書化された病理を表す。{27}"
        ),
        "治療的命名と疾患喧伝の区別は三つの次元に沿って引くことができると論じる（表2）。",
    ]
    for text in boundary_ja:
        p = doc.add_paragraph()
        if isinstance(text, tuple):
            add_superscript_text(p, text[0] if isinstance(text, tuple) else text)
        else:
            add_superscript_text(p, text)

    doc.add_paragraph()

    # ── 表2 ──
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    run = p.add_run("表2. ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run("治療的命名と疾患喧伝の区別")
    run.font.size = Pt(10)

    table2_ja = [
        ["次元", "治療的命名", "疾患喧伝"],
        ["起源", "臨床的必要性と患者の苦痛が駆動；臨床専門家と患者擁護者のコンセンサスにより開発",
         "商業的利益が駆動；製薬マーケティングと後援された疾患啓発キャンペーンにより開発"],
        ["エビデンス基盤", "明確な症状パターン、メカニズム、治療反応の経験的観察に基礎づけ；フィールドテストで検証",
         "正常変異の医療化またはリスク因子の疾患再定義に基づく；独立した検証が限定的"],
        ["アウトカム軌道", "患者アウトカムの改善、資源配分の改善、時間経過による理解の精緻化（NTL正のサイクル）",
         "過剰診断、不必要な治療、医原性傷害、資源の不適切配分"],
    ]

    table2 = doc.add_table(rows=len(table2_ja), cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(table2_ja):
        row = table2.rows[i]
        for j, cell_text in enumerate(row_data):
            set_cell_text(row.cells[j], cell_text, bold=(i == 0), size=9)

    doc.add_paragraph()

    boundary_cont_ja = [
        (
            "ICD-11慢性疼痛分類は三つの基準すべてにおいて治療的命名を例証する。文書化された臨床的"
            "必要性に応じて国際臨床専門家タスクフォースにより開発され、{17} WHOフィールドテストで"
            "検証され、{18} 初期のエビデンスは診断精度と臨床的有用性の改善を示唆している。{16}"
        ),
        (
            "さらなる境界条件はKleinmanの観察により提供される。特定の疾患経験はその文化的命名とは"
            "独立に存続する。{12} 拒食症は痩身が評価されなかった時代にも存在し、慢性疼痛はそれが"
            "命名されるか否かにかかわらず存続する。医療版サピア＝ウォーフ仮説は命名が疾患を創出"
            "すると主張するのではない。命名が疾患への臨床的対応を形作ると主張するのである。"
        ),
    ]
    for text in boundary_cont_ja:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 7. 研究アジェンダ
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("7. 提案する研究アジェンダ")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    agenda_ja = [
        (
            "ICD-10からICD-11への移行は、医療版サピア＝ウォーフ仮説を実証的に検証するために"
            "活用可能な治療的命名の自然実験を構成する。以下の研究アジェンダを提案する："
        ),
    ]
    for text in agenda_ja:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    research_ja = [
        ("前後比較診断パターン分析: ",
         "ICD-11を採用した医療システムの行政データを用いて移行前後の慢性疼痛診断パターンを比較。"),
        ("患者報告アウトカム研究: ",
         "同等の臨床的プレゼンテーションに対してICD-10 vs ICD-11カテゴリーで診断された患者の満足度、"
         "診断確実性、アドヒアランス、自己報告疼痛アウトカムを比較する前向きコホート研究。"),
        ("臨床家認知実験: ",
         "ICD-11慢性疼痛カテゴリーの研修ありなしの臨床家に同一の疼痛プレゼンテーションを提示する"
         "ランダム化ビニエット研究。"),
        ("研究資金分析: ",
         "ICD-11採用前後の慢性疼痛研究出版物と資金配分の計量書誌学的分析。"),
        ("国際比較: ",
         "国ごとのICD-11採用時期のずれを利用した差の差分析デザインで慢性疼痛アウトカムを比較。"),
    ]

    for i, (bold_part, normal_part) in enumerate(research_ja, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(f"({i}) ")
        run.font.size = Pt(11)
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(normal_part)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 8. 結論
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("8. 結論")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    conclusion_ja = [
        (
            "サピア＝ウォーフ仮説は、言語が現実の受動的な鏡ではなくその構築への能動的参加者で"
            "あることを教える。医療版サピア＝ウォーフ仮説はこの洞察を臨床医学に拡張する。"
            "疾患命名体系は単なるラベルではなく、知覚を形作り、注意を向け、資源を配分し、"
            "最終的に患者アウトカムに影響を与える認知的・制度的ツールである。"
        ),
        (
            "疼痛医学はこの命題にとって最も説得力のある領域を提供する。ICD-10における慢性疼痛の"
            "歴史的分類不足は単なる行政的不便ではなく、苦痛の構造的原因であった。ICD-11慢性疼痛"
            "分類、nociplastic painの概念、慢性一次性疼痛の疾患としての認定は、治療的命名の行為を"
            "集合的に代表する。それらは臨床的現実を記述するだけでなく、能動的に改善する。"
        ),
        (
            "疾患命名治療ループは、命名がいかに治療的効果を生み出すかを理解し、治療的命名と"
            "疾患喧伝を区別する境界条件を同定するための形式的枠組みを提供する。"
        ),
        (
            "医療版サピア＝ウォーフ仮説が正しければ、疼痛医学における疾患命名カテゴリーの"
            "意図的かつエビデンスに基づく拡張は、単なる分類の官僚的行為ではなく治療介入である。"
            "その含意は実践的かつ深遠である。疼痛医学において、疾患名の精度と豊かさを高める"
            "ことは、利用可能な最もコスト効果の高い治療戦略の一つとなりうる。"
        ),
    ]
    for text in conclusion_ja:
        p = doc.add_paragraph()
        add_superscript_text(p, text)

    doc.add_paragraph()

    # ══════════════════════════════════════════════
    # 利益相反・資金・文献
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("利益相反")
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("著者らは利益相反がないことを宣言する。")
    run.font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("資金")
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("[資金情報追記予定]")
    run.font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("文献")
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph()

    # Same references as English version (Vancouver style)
    references = [
        "Sapir E. The status of linguistics as a science. Language. 1929;5(4):207-214.",
        "Whorf BL. Language, Thought, and Reality: Selected Writings. Cambridge, MA: MIT Press; 1956.",
        "Winawer J, Witthoft N, Frank MC, et al. Russian blues reveal effects of language on color discrimination. Proc Natl Acad Sci U S A. 2007;104(19):7780-7785.",
        "Levinson SC. Frames of reference and Molyneux's question: cross-linguistic evidence. In: Bloom P, et al., eds. Language and Space. Cambridge, MA: MIT Press; 1996:109-169.",
        "Kosek E, Cohen M, Baron R, et al. Do we need a third mechanistic descriptor for chronic pain states? Pain. 2016;157(7):1382-1386.",
        "Jutel A. Putting a Name to It: Diagnosis in Contemporary Society. Baltimore: Johns Hopkins University Press; 2011.",
        "Bontempo AC. Does labeling a patient with fibromyalgia alter health care utilization? J Gen Intern Med. 2021;36(5):1352-1358.",
        "Kaplan RM. Disease, Diagnoses, and Dollars. New York: Copernicus Books; 2009.",
        "World Health Organization. International Classification of Diseases 11th Revision. Geneva: WHO; 2019.",
        "Hacking I. The looping effects of human kinds. In: Sperber D, et al., eds. Causal Cognition. Oxford: Clarendon Press; 1995:351-394.",
        "Tsou JY. Hacking on the looping effects of psychiatric classifications. Int Stud Philos Sci. 2007;21(3):329-344.",
        "Kleinman A. The Illness Narratives: Suffering, Healing, and the Human Condition. New York: Basic Books; 1988.",
        "Breivik H, Collett B, Ventafridda V, et al. Survey of chronic pain in Europe. Eur J Pain. 2006;10(4):287-333.",
        "GBD 2021 Diseases and Injuries Collaborators. Global incidence, prevalence, and years lived with disability for 354 diseases and injuries. Lancet. 2024;403(10440):2162-2199.",
        "Rief W, Kaasa S, Jensen R, et al. The need to revise pain diagnoses in ICD-10. Pain. 2010;149(2):169-170.",
        "Nugraha B, Gutenbrunner C, Barke A, et al. The IASP classification of chronic pain for ICD-11: functioning properties. Pain. 2019;160(1):88-94.",
        "Treede RD, Rief W, Barke A, et al. Chronic pain as a symptom or a disease: the IASP Classification for ICD-11. Pain. 2019;160(1):19-27.",
        "Barke A, Korwisi B, Jakob R, et al. Classification of chronic pain for ICD-11: results of WHO field testing. Pain. 2022;163(2):e285-e292.",
        "Kosek E, Clauw D, Nijs J, et al. Chronic nociplastic pain affecting the musculoskeletal system: clinical criteria. Pain. 2021;162(11):2629-2634.",
        "Nijs J, Malfliet A, Nishigami T. Nociplastic pain and central sensitization: a terminology update. Braz J Phys Ther. 2023;27(3):100518.",
        "Fitzcharles MA, Cohen SP, Clauw DJ, et al. Nociplastic pain: towards an understanding of prevalent pain conditions. Lancet. 2021;397(10289):2098-2110.",
        "O'Keeffe M, Ferreira GE, Harris IA, et al. Effect of diagnostic labelling on management intentions for non-specific low back pain. Eur J Pain. 2022;26(7):1532-1545.",
        "Zadro JR, Ferreira GE, O'Keeffe M, et al. The effect of diagnostic labels on treatment preferences and beliefs. Arch Phys Med Rehabil. 2025 (in press).",
        "Sari DM, Samuels K, Engel L, et al. How patients with an uncertain diagnosis experience intolerance of uncertainty. Psychol Res Behav Manag. 2021;14:1269-1281.",
        "Tanna V, Heathcote LC, Heirich MS, et al. Diagnostic uncertainty in children with chronic pain. Children. 2020;7(10):165.",
        "O'Sullivan E, Shanahan E, Costelloe L, et al. Importance of diagnosis in chronic pain: patient's perspective. Pain Stud Treat. 2026;14(1):1-12.",
        "Moynihan R, Heath I, Henry D. Selling sickness: the pharmaceutical industry and disease mongering. BMJ. 2002;324(7342):886-891.",
        "Moynihan R, Cassels A. Selling Sickness. New York: Nation Books; 2005.",
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.font.size = Pt(10)
        run = p.add_run(ref)
        run.font.size = Pt(10)

    # ── Save ──
    output_path = os.path.join(OUTPUT_DIR, "manuscript_ja.docx")
    doc.save(output_path)
    print(f"Manuscript saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_manuscript()
